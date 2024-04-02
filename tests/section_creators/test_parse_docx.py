import os
from docx import Document
from section_creators.docx_parser import parse_docx
from tests.base_test import BaseTest


class TestParseDocx(BaseTest):

    def setUp(self):
        super().setUp()
        # Create a temporary Word document for testing purposes
        self.test_docx_path = 'test.docx'
        doc = Document()
        doc.add_heading('Header 1', level=1)
        doc.add_paragraph('This is a test paragraph under Header 1.')
        doc.add_heading('Header 2', level=1)
        doc.add_paragraph('This is a test paragraph under Header 2.')
        doc.add_paragraph('This is a second test paragraph under Header 2.')
        doc.save(self.test_docx_path)

    def test_parse_docx(self):
        empty_anchor_url = ''
        expected_result = [
            ('Header 1', 'This is a test paragraph under Header 1.', empty_anchor_url),
            ('Header 2', 'This is a test paragraph under Header 2.This is a second test paragraph under Header 2.', empty_anchor_url),
        ]

        result = list(parse_docx(self.test_docx_path))
        self.assertEqual(result, expected_result, 'The parsed docx content does not match the expected result.')

    def tearDown(self):
        # Remove the temporary Word document after testing
        os.remove(self.test_docx_path)
